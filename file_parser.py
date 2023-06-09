import os
import re
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import numpy as np
import string
import win32com.client

def remove_non_ascii(text):
    ascii_chars = set(string.printable)
    return ''.join(
        filter(lambda x: x in ascii_chars, text)
    )

class PdfParser:
    def __init__(self, file_path, temp_path):
        self.email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        self.temp_path = temp_path
        self.file_path = file_path

    def extruct_mail(self):
        # Init variables
        email_list = []

        # Convert PDF to images
        images = convert_from_path(self.file_path, output_folder=self.temp_path)

        # Loop through each image and perform OCR
        for _, image in enumerate(images):
            # Perform OCR on the image using pytesseract
            text = pytesseract.image_to_string(image)
            filtered_text = remove_non_ascii(text)
            # Find mail address in the text
            email_matches = re.findall(self.email_regex, filtered_text)
            if email_matches == []:
                pass
            email_list += email_matches

        return email_list


class WordParser:
    def __init__(self, file_path):
        self.email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        file_type = os.path.splitext(file_path)[1]
        if file_type == '.docx':
            self.doc = Document(file_path)
        else:
            self.doc = self.convert_doc_to_docx(file_path)

    @staticmethod
    def convert_doc_to_docx(file_path):
        new_file_path = os.path.splitext(file_path)[0] + '.docx'

        wrd = win32com.client.Dispatch("Word.Application")
        wrd.visible = 0
        wb = wrd.Documents.Open(file_path)
        wb.SaveAs2(new_file_path, FileFormat=12)
        wb.Close()
        wrd.Quit()

        return Document(new_file_path)


    def extruct_mail(self):
        # Init Variables
        email_list = []

        # Extruct from HyperLnk
        email_list += self.extruct_from_hyperlink(self.doc.part)

        # Extruct Text from Paragraphs
        for paragraph in self.doc.paragraphs:
            email_matches = self.extruct_from_paragraph(paragraph)
            email_list += email_matches

        # Extruct Text from sections
        for section in self.doc.sections:
            # Extruct Text from header
            header = section.header
            email_list += self.extruct_from_hyperlink(header.part)
            for paragraph in header.paragraphs:
                email_matches = self.extruct_from_paragraph(paragraph)
                email_list += email_matches

            # Extruct Text from footer
            footer = section.footer
            email_list += self.extruct_from_hyperlink(footer.part)
            for paragraph in footer.paragraphs:
                email_matches = self.extruct_from_paragraph(paragraph)
                email_list += email_matches

        # # Extruct Text from Text Box:
        # for shape in self.doc.inline_shapes:
        #     if shape.has_text_frame:
        #         text_frame = shape.text_frame
        #         for paragraph in text_frame.paragraphs:
        #             email_matches = self.extruct_from_paragraph(paragraph)
        #             email_list += email_matches

        # Extruct Text from tables:
        for tbl in self.doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        email_matches = self.extruct_from_paragraph(paragraph)
                        email_list += email_matches

        return email_list

    def extruct_from_paragraph(self, paragraph):
        email_matches = []
        text = paragraph.text
        if len(text) > 0:
            email_matches += re.findall(self.email_regex, text)

        return email_matches

    def extruct_from_hyperlink(self, part):
        email_list = []
        rels = part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                email_list += re.findall(self.email_regex, rels[rel]._target)

        return email_list


def filter_list(email_list):
    return np.unique(email_list).tolist()


def test():
    file_path = r'C:\temp\mail_handler\test.docx'
    parser = WordParser(file_path)
    mail = parser.extruct_mail()
    filtered_mail = filter_list(mail)
    print('; '.join(filtered_mail))

    file_path = r'C:\temp\mail_handler\נבו_חיון.pdf'
    parser = PdfParser(file_path, r'C:\temp\mail_handler')
    mail = parser.extruct_mail()
    filtered_mail = filter_list(mail)
    print('; '.join(filtered_mail))


if __name__ == '__main__':
    test()
